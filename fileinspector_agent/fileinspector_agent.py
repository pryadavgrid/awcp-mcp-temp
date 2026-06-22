"""An autonomous, governed FILE-INSPECTOR WORKER runtime (AWCP agent-on-a-runtime).

Give it ANY file — text, image, PDF, JSON, .docx, CSV, source code, or an unknown
binary — and it tells you WHAT the file is and WHAT is inside it. It pulls GOALS
off the task queue (each goal carries a local FILE_PATH, set by the upload box or
typed by hand) and executes each in multiple steps:
  - classify_file  : detect type + structural metadata (deterministic, never fails)
  - read_document  : extract text from PDF / DOCX / TXT / JSON / CSV / code
  - describe_image : caption an image with a local Ollama VISION model (graceful
                     fallback to rich metadata if no vision model is installed)
  - save_artifact  -> governed LOCAL write  (medium risk, gated)
  - external_post   -> governed EXTERNAL write (high risk, gated + needs approval)
Queue/worker/governance/approval/UI live in awcp_kit; this file supplies the
LangGraph framework agent + the run_goal() hook.

Run as:  python agent_runtime.py   (absolute path via run.sh so the detector sees
the `langgraph` import).
"""

import base64
import io
import json
import os
import re

from langgraph.graph import StateGraph  # noqa: F401  (import marks this as LangGraph)
from langchain_core.tools import tool
from langchain_core.messages import HumanMessage, SystemMessage
from langchain_ollama import ChatOllama

from fastapi import FastAPI
import uvicorn

import awcp_kit as kit

MODEL = os.getenv("FILE_MODEL", "qwen2.5:7b")           # text reading / summary
VISION_MODEL = os.getenv("VISION_MODEL", "qwen2.5vl:7b")  # image captioning
OLLAMA_BASE = os.getenv("OLLAMA_BASE", "http://localhost:11434")
PORT = int(os.getenv("FILE_PORT", "8104"))
HERE = os.path.dirname(os.path.abspath(__file__))

MAX_TEXT_CHARS = int(os.getenv("FILE_MAX_CHARS", "40000"))
MODEL_CONTEXT_CHARS = int(os.getenv("FILE_MODEL_CHARS", "16000"))  # cap fed to the model

# Extension -> coarse file class. Anything unlisted falls back to magic-byte sniffing.
IMAGE_EXT = {".jpg", ".jpeg", ".png", ".webp", ".bmp", ".gif", ".tiff", ".tif"}
TEXT_EXT = {".txt", ".md", ".json", ".csv", ".tsv", ".py", ".js", ".ts", ".html",
            ".htm", ".xml", ".ini", ".cfg", ".yaml", ".yml", ".toml", ".log",
            ".sql", ".sh", ".java", ".c", ".cpp", ".go", ".rs", ".rb"}

# First bytes -> human label, for files whose extension tells us nothing.
MAGIC = [
    (b"%PDF", "PDF document"),
    (b"\x89PNG", "PNG image"),
    (b"\xff\xd8\xff", "JPEG image"),
    (b"GIF8", "GIF image"),
    (b"PK\x03\x04", "ZIP archive / OOXML (docx, xlsx, …)"),
    (b"\x1f\x8b", "GZIP archive"),
    (b"\x7fELF", "ELF executable"),
    (b"MZ", "Windows PE executable"),
    (b"\xca\xfe\xba\xbe", "Java class / Mach-O fat binary"),
]

SYSTEM = (
    "You are a FILE-INSPECTION assistant. You are handed the EXTRACTED CONTENT of a "
    "file and must tell the user, in a BRIEF plain-language summary (2-5 sentences), "
    "what the file is about and what is inside it. Describe the actual contents and "
    "their meaning. Do NOT mention technical metadata — no file size, byte counts, "
    "dimensions, formats, encodings, page counts, or character counts. Use ONLY the "
    "extracted content as your source of truth; never invent details, and never emit "
    "tool-call syntax or JSON. Just write the short summary."
)


# --------------------------------------------------------------------------
# Deterministic inspection helpers (no model needed — these always work).
# --------------------------------------------------------------------------
def _classify(path: str) -> dict:
    if not os.path.exists(path):
        return {"error": f"file not found: {path}"}
    filename = os.path.basename(path)
    size = os.path.getsize(path)
    ext = os.path.splitext(filename)[1].lower()
    info = {"filename": filename, "size_bytes": size, "extension": ext or "(none)"}

    if ext in IMAGE_EXT:
        try:
            from PIL import Image
            with Image.open(path) as im:
                im.verify()
            with Image.open(path) as im:
                info.update(file_type="image", width=im.width, height=im.height,
                            image_format=im.format, mode=im.mode)
        except Exception as e:  # noqa: BLE001
            info.update(file_type="image", warning=f"unreadable image: {e}")
        return info

    if ext == ".pdf":
        try:
            import pypdf
            reader = pypdf.PdfReader(path)
            info.update(file_type="pdf", pages=len(reader.pages),
                        is_encrypted=reader.is_encrypted)
        except Exception as e:  # noqa: BLE001
            info.update(file_type="pdf", warning=f"unreadable pdf: {e}")
        return info

    if ext == ".docx":
        try:
            import docx
            doc = docx.Document(path)
            info.update(file_type="docx", paragraphs=len(doc.paragraphs),
                        tables=len(doc.tables))
        except Exception as e:  # noqa: BLE001
            info.update(file_type="docx", warning=f"unreadable docx: {e}")
        return info

    if ext in TEXT_EXT:
        info["file_type"] = "text"
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                snippet = f.read(2048)
            info["character_sample"] = len(snippet)
            info["lines_sampled"] = snippet.count("\n") + 1
            if ext == ".json":
                with open(path, "r", encoding="utf-8", errors="ignore") as f:
                    json.load(f)
                info["json"] = "valid"
        except json.JSONDecodeError as e:
            info["json"] = f"invalid JSON: {e}"
        except Exception as e:  # noqa: BLE001
            info["warning"] = str(e)
        return info

    # Unknown extension -> sniff magic bytes.
    try:
        with open(path, "rb") as f:
            head = f.read(16)
    except Exception as e:  # noqa: BLE001
        return {**info, "file_type": "binary", "warning": str(e)}
    label = next((lbl for sig, lbl in MAGIC if head.startswith(sig)), None)
    # Heuristic: mostly-printable bytes => treat as text.
    if label is None and head and all(b in b"\t\n\r\f" or 32 <= b < 127 for b in head):
        info["file_type"] = "text"
        return info
    info.update(file_type="binary", guessed=label or "unknown binary format",
                hex_signature=head.hex().upper())
    return info


def _extract_text(path: str, file_type: str) -> str:
    if file_type == "pdf":
        import pypdf
        reader = pypdf.PdfReader(path)
        out = []
        for i, page in enumerate(reader.pages):
            t = page.extract_text() or ""
            if t.strip():
                out.append(f"--- Page {i + 1} ---\n{t}")
        return "\n".join(out)
    if file_type == "docx":
        import docx
        doc = docx.Document(path)
        out = [p.text for p in doc.paragraphs if p.text.strip()]
        for ti, table in enumerate(doc.tables):
            out.append(f"\n--- Table {ti + 1} ---")
            for row in table.rows:
                out.append(" | ".join(c.text.strip() for c in row.cells))
        return "\n".join(out)
    # plain text / json / csv / code
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def _truncate(text: str) -> str:
    if len(text) > MAX_TEXT_CHARS:
        return text[:MAX_TEXT_CHARS] + "\n...[truncated]..."
    return text


# --------------------------------------------------------------------------
# Tools (read tools run freely; writes are governed by awcp_kit).
# --------------------------------------------------------------------------
@tool
def classify_file(path: str) -> str:
    """Classify a file by path: returns its type (image/pdf/docx/text/binary) and
    structural metadata (size, dimensions, page count, JSON validity, magic bytes).
    Call this FIRST. Deterministic — does not read the full contents."""
    return json.dumps(_classify(path), indent=2)


@tool
def read_document(path: str) -> str:
    """Extract and return the readable TEXT content of a document (PDF, .docx, .txt,
    .json, .csv, source code, etc.). Use for non-image files after classify_file."""
    info = _classify(path)
    if "error" in info:
        return info["error"]
    ft = info.get("file_type", "text")
    if ft == "image":
        return "This is an image, not a text document. Use describe_image instead."
    try:
        text = _extract_text(path, ft)
    except Exception as e:  # noqa: BLE001
        return f"Could not extract text: {e}"
    if not text.strip():
        return f"No extractable text found in this {ft} file (it may be empty, scanned, or binary)."
    return _truncate(text)


@tool
def describe_image(path: str,
                   question: str = "Briefly describe what is shown in this image in 2-4 sentences.") -> str:
    """Look at an image and briefly describe what it shows, using a local Ollama
    vision model."""
    info = _classify(path)
    if "error" in info:
        return info["error"]
    try:
        from PIL import Image
        im = Image.open(path)
        if im.mode != "RGB":
            im = im.convert("RGB")
        im.thumbnail((1024, 1024), Image.Resampling.LANCZOS)
        buf = io.BytesIO()
        im.save(buf, format="JPEG", quality=85)
        b64 = base64.b64encode(buf.getvalue()).decode()
    except Exception as e:  # noqa: BLE001
        return f"Could not load image: {e}"
    try:
        vlm = ChatOllama(model=VISION_MODEL, base_url=OLLAMA_BASE, temperature=0.2)
        msg = HumanMessage(content=[
            {"type": "text", "text": question},
            {"type": "image_url", "image_url": f"data:image/jpeg;base64,{b64}"},
        ])
        return vlm.invoke([msg]).content
    except Exception as e:  # noqa: BLE001
        return (f"Could not describe this image — the vision model '{VISION_MODEL}' "
                f"is unavailable ({e}). Install it with `ollama pull {VISION_MODEL}`.")


@tool
def save_artifact(name: str, content: str) -> str:
    """Save an inspection report to disk. GOVERNED local write (gated)."""
    return kit.save_artifact(name, content)


@tool
def external_post(summary: str) -> str:
    """Submit/publish an inspection report to an external system. HIGH-RISK governed
    write: gated AND pauses for operator approval. Use only when the goal asks to
    submit, send, publish, or report the findings externally."""
    return kit.external_post(summary)


TOOLS = [classify_file, read_document, describe_image, save_artifact, external_post]
TOOL_NAMES = [t.name for t in TOOLS]

_llm = ChatOllama(model=MODEL, base_url=OLLAMA_BASE, temperature=0)


def _extract_path(goal: str):
    """Pull the local file path out of a goal. The upload box appends
    'FILE_PATH: <path>'; a hand-typed goal may just mention an existing path."""
    m = re.search(r"FILE_PATH:\s*(.+)", goal)
    if m:
        p = m.group(1).strip().strip('"').strip("'")
        if os.path.isfile(p):
            return p
    for tok in re.findall(r"\"[^\"]+\"|'[^']+'|\S+", goal):
        t = tok.strip("\"'")
        if os.path.isfile(t):
            return t
    return None


def _looks_like_junk(answer: str) -> bool:
    """Detect weak-model failures: empty output, or a tool call printed as text."""
    a = (answer or "").strip()
    if len(a) < 15:
        return True
    return bool(re.search(r'"(name|parameters|arguments)"\s*:', a))


def _brief_fallback(content: str) -> str:
    """Used only when the summariser model is unavailable: show the file's content
    directly (no metadata), so the user still sees what's inside."""
    body = content.strip()
    return ("_(Summary model unavailable — showing the file's content directly.)_\n\n"
            + (body[:2000] + (" …" if len(body) > 2000 else "")))


def run_goal(goal: str) -> dict:
    """Framework hook: inspect the file referenced by the goal and return a BRIEF
    summary of what's inside it (no metadata). Images are summarised by the vision
    model; everything else is read deterministically, then summarised by the text
    model. Governed writes happen via the kit's finalize step."""
    path = _extract_path(goal)
    if not path:
        return {"result": "No file found to inspect. Attach a file (or include a "
                          "valid FILE_PATH) and I'll tell you what's inside it.",
                "tools_used": []}

    info = _classify(path)
    if "error" in info:
        return {"result": info["error"], "tools_used": ["classify_file"]}

    ft = info.get("file_type", "text")
    request = re.sub(r"FILE_PATH:.*", "", goal, flags=re.S).strip()
    tools = ["classify_file"]

    # Images: the vision model both "reads" and summarises in one step.
    if ft == "image":
        question = request or "Briefly describe what is shown in this image in 2-4 sentences."
        tools.append("describe_image")
        return {"result": describe_image.func(path, question), "tools_used": tools}

    # Documents / text: extract deterministically, then summarise with the text model.
    if ft in ("pdf", "docx", "text"):
        content = read_document.func(path)
        tools.append("read_document")
    else:  # binary — nothing readable to summarise
        return {"result": "This looks like a binary file with no human-readable text "
                          "content to summarise.", "tools_used": tools}

    instruction = request or "Give a brief summary of what this file contains."
    prompt = f"EXTRACTED CONTENT:\n{content[:MODEL_CONTEXT_CHARS]}\n\nTASK: {instruction}"
    try:
        _resp = _llm.invoke([SystemMessage(content=SYSTEM),
                             HumanMessage(content=prompt)])
        answer = _resp.content
        if _looks_like_junk(answer):
            answer = _brief_fallback(content)
        _um = getattr(_resp, "usage_metadata", None) or {}
        _usage = {"input_tokens": int(_um.get("input_tokens", 0) or 0),
                  "output_tokens": int(_um.get("output_tokens", 0) or 0)}
        return {"result": answer, "tools_used": tools, "usage": _usage}
    except Exception as e:  # noqa: BLE001  (e.g. Ollama not running) -> still answer
        return {"result": _brief_fallback(content), "tools_used": tools, "error": str(e)}


app = FastAPI(title="File Inspector Worker Runtime")

if __name__ == "__main__":
    kit.mount(
        app,
        meta={"agent": "File Inspector", "framework": "langgraph",
              "model": MODEL, "tools": TOOL_NAMES, "dir": HERE,
              "purpose": "Universal file inspector — identify any file (text, image, PDF, JSON, .docx, code, binary) and explain what's inside.",
              "format": "markdown", "accent": "#f59e0b", "logo": "\U0001F5C2",
              "accepts_files": True,
              "examples": ["Identify this file and summarise its contents.",
                           "What type of file is this and what does it contain?",
                           "Extract the key facts from this document.",
                           "What is shown in this image?"]},
        run_goal=run_goal,
    )
    print(f"\U0001F5C2  File Inspector WORKER  →  http://localhost:{PORT}   (model={MODEL}, vision={VISION_MODEL})")
    uvicorn.run(app, host="0.0.0.0", port=PORT)
